#Resume Adapter for Individual Applications: python 3.11.2 64-bit

import docx
import re
from dateutil import parser
import nltk
from nltk.tokenize import word_tokenize
from nltk.corpus import stopwords
import string
doc = docx.Document('rkresume.docx')


import requests
from bs4 import BeautifulSoup

# Send an HTTP request to the URL of the webpage you want to access
url = "https://www.jobscan.co/blog/top-resume-keywords-boost-resume/#customer-service-retail-human-resources-tourism"
response = requests.get(url)

# Parse the HTML content of the page using BeautifulSoup
soup = BeautifulSoup(response.content, 'html.parser')

# Find the section of the webpage that contains the skills
skill_section = soup.find('li', {'class': 'li'})

# Find all the skills in the skill section
skills = [skill.text for skill in skill_section.find_all('strong')]

# Print the list of skills
print(skills)
exit()
url = "https://pace.joinhandshake.com/stu/jobs/6655814?ref=open-in-new-tab&search_id=bf671888-a714-4f09-ae0b-3f47d05d5267"

def contains_month(text_list):
    for i, text in enumerate(text_list):
        if(text[0].isdigit() == False):
            try:
                dt = parser.parse(text)
                month_index = dt.month - 1  # month is 1-based, so subtract 1 to get the index
                return i
            except ValueError:
                pass  # ignore non-date strings
    return -1


def strip_parts(text):
    tokens = word_tokenize(text)

    punctuations = string.punctuation
    # Get the English stop words list from NLTK
    stop_words = set(stopwords.words('english'))

    # Remove stop words from the list of tokens
    filtered_tokens = [word for word in tokens if not word.lower() in stop_words and not word in punctuations and all(char.isalpha() for char in word) ]

    # Join the filtered tokens back into a string
    filtered_text = ' '.join(filtered_tokens)

    # Get the POS tags for the words
    pos_tags = nltk.pos_tag(filtered_tokens)

    # Print the POS tags
    lexicon = []
    for tag in pos_tags:
        lexicon.append(tag[0])
    return lexicon

'''
POSTINGS:

https://pace.joinhandshake.com/stu/jobs/6655814?ref=open-in-new-tab&search_id=bf671888-a714-4f09-ae0b-3f47d05d5267
https://pace.joinhandshake.com/stu/jobs/7432936?ref=open-in-new-tab&search_id=fb5a8b32-8c53-4d3b-aa3f-b5c3de4ec4e2
https://pace.joinhandshake.com/stu/jobs/7608071?ref=open-in-new-tab&search_id=bf671888-a714-4f09-ae0b-3f47d05d5267
https://www.indeed.com/jobs?q=software+engineer+intern+summer+2023&l=New+York%2C+NY&from=searchOnHP&vjk=cd63f0986d6fd807

'''
def generate_better_prompt(start_text, text):
    new_text = ""
    

    # Print the extracted text
    print(text)
    return new_text
replaceables = {}
tag = ""
i = 0
idx = 0
for para in doc.paragraphs:
    if(para.text == "PROFILE"):
            replaceables.update({"profile":i})
    elif(para.text == "PROJECTS"):
        idx = i
    if(idx > 0 and i > idx):
        t = re.sub(r'\s+', ' ', para.text).split()
        c = contains_month(t)
        if(c != -1):
            replaceables.update({" ".join(t[0:c]):i})
            
    i+= 1
    # print(para.text)
print(replaceables)

text = ""
with open('input.txt', 'r') as file:
    text = file.read()

lex = strip_parts(text)


for key in replaceables.keys():
    idx = int(replaceables.get(key)) + 1
    target_paragraph = doc.paragraphs[idx]  

    paragraph_format = target_paragraph.paragraph_format
    new_text = generate_better_prompt(target_paragraph.text ,lex)
    print(new_text)
    exit()
    target_paragraph.text = new_text

    target_paragraph.paragraph_format = paragraph_format

    print(doc.paragraphs[int(replaceables.get(key)) + 1].text)

doc.save('Kroening, Ronald- Resume.docx')